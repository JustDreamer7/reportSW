import os
from datetime import date

import pandas as pd
from cycler import cycler
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
from docx.shared import Cm
from docx.shared import Inches
from docx.shared import Pt
# from docx.shared import RGBColor
from matplotlib import pyplot as plt

from a4fr import a4fr
from a52det import a52det
from ntozeromas import ntozerotr
from timestop import timeBreak
from timework import timeWork
from ntozeromas import change_name
from corr_a52det import main, make_event


# функции для проставления номера страницы в ворде -------------------------------
def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):
    # выравниваем параграф по центру
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # запускаем динамическое обновление параграфа
    page_num_run = paragraph.add_run()
    # обозначаем начало позиции вывода
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')
    # задаем вывод текущего значения страницы PAGE (всего страниц NUMPAGES)
    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"
    # обозначаем конец позиции вывода
    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')
    # добавляем все в наш параграф (который формируется динамически)
    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)


# -------------------------------------------------------------------------------------------------------


# Длиннокод того, как строяться графики и заполняется ворд
def secProccesing(stday, stmonth, styear, endday, endmonth, endyear, path, pathpic, file1cl, file2cl):
    amp = 6
    fr = 2
    amp_2 = 11
    fr_2 = 1
    datedirect = pathpic + '/{}'.format(styear)
    if ~os.path.exists(datedirect):
        try:
            os.mkdir(datedirect)
        except OSError:
            print("Создать директорию %s не удалось" % datedirect)
        else:
            print("Успешно создана директория %s " % datedirect)
    your_format = lambda x: '{{:0.{}f}}'.format(2 if x > 1 else 3).format(x)
    pd.set_option('display.float_format', your_format)
    a = date(styear, stmonth, stday)
    b = date(endyear, endmonth, endday)
    worktime = timeWork(1, stday, endday, styear, endyear, stmonth, endmonth, file1cl)
    worktime_2 = timeWork(2, stday, endday, styear, endyear, stmonth, endmonth, file2cl)
    # print(worktime)
    # print(worktime_2)

    timestop = timeBreak(1, stday, endday, styear, endyear, stmonth, endmonth, file1cl)
    timestop_2 = timeBreak(2, stday, endday, styear, endyear, stmonth, endmonth, file2cl)

    print(timestop_2)
    print(timestop)
    both = pd.DataFrame()
    for i in range(len(timestop.index)):
        # j: int
        for j in range(len(timestop_2.index)):
            if timestop['DATE'][i] == timestop_2['DATE'][j]:
                SM1 = timestop.iloc[i][1] * 60 + timestop.iloc[i][2]
                EM1 = timestop.iloc[i][3] * 60 + timestop.iloc[i][4]
                SM2 = timestop_2.iloc[j][1] * 60 + timestop_2.iloc[j][2]
                EM2 = timestop_2.iloc[j][3] * 60 + timestop_2.iloc[j][4]
                ser = pd.Series([SM1, EM1, SM2, EM2])
                both[i] = ser.copy()
                # строка выше - тест
    if len(both.index) > 0:
        both.index = ["SM1", "EM1", "SM2", "EM2"]
    bothdt = both.T
    bothtime = 0
    print(bothdt)
    if len(bothdt) > 0:
        start_minutes_1 = bothdt['SM1'].tolist()
        start_minutes_2 = bothdt['SM2'].tolist()
        end_minutes_1 = bothdt['EM1'].tolist()
        end_minutes_2 = bothdt['EM2'].tolist()
        for i in range(len(bothdt.index)):
            if start_minutes_1[i] <= start_minutes_2[i] < end_minutes_1[i]:
                bothtime += min(end_minutes_2[i], end_minutes_1[i]) - max(start_minutes_2[i], start_minutes_1[i])
            elif start_minutes_2[i] <= start_minutes_1[i] < end_minutes_2[i]:
                bothtime += min(end_minutes_2[i], end_minutes_1[i]) - max(start_minutes_2[i], start_minutes_1[i])

    # Обернуть в функцию

    breaks = len(timestop.index)
    failstr_begin = []
    failstr_end = []
    lost_minutes = []
    for i in range(len(timestop.index)):
        failstr_begin.append(
            " {}  {:02}:{:02}".format(timestop['DATE'][i].date(), timestop['stHour'][i], timestop['stMinutes'][i]))
        failstr_end.append(
            " {}  {:02}:{:02}".format(timestop['DATE'][i].date(), timestop['endHour'][i], timestop['endMinutes'][i]))
        lost_minutes.append(timestop['endHour'][i] * 60 + timestop['endMinutes'][i] - (
                timestop['stHour'][i] * 60 + timestop['stMinutes'][i]))

    for i in range(1, len(timestop.index)):
        if timestop['stHour'][i] == 0 and timestop['stMinutes'][i] == 5 and timestop['endHour'][i - 1] == 24 and \
                (timestop['DATE'][1] - timestop['DATE'][0]).days == 1:
            breaks -= 1

    breaks_2 = len(timestop_2.index)
    failstr_2_begin = []
    failstr_2_end = []
    lost_minutes_2 = []
    for i in range(len(timestop_2.index)):
        failstr_2_begin.append(
            " {}  {:02}:{:02}".format(timestop_2['DATE'][i].date(), timestop_2['stHour'][i],
                                      timestop_2['stMinutes'][i]))
        failstr_2_end.append(
            " {}  {:02}:{:02}".format(timestop_2['DATE'][i].date(), timestop_2['endHour'][i],
                                      timestop_2['endMinutes'][i]))
        lost_minutes_2.append(timestop_2['endHour'][i] * 60 + timestop_2['endMinutes'][i] - (
                timestop_2['stHour'][i] * 60 + timestop_2['stMinutes'][i]))

    for i in range(1, len(timestop_2.index)):
        if timestop_2['stHour'][i] == 0 and timestop_2['stMinutes'][i] == 5 and timestop_2['endHour'][i - 1] == 24 and \
                (timestop_2['DATE'][1] - timestop_2['DATE'][0]).days == 1:
            breaks_2 -= 1

    realtime = worktime_2['WORKTIME'].sum() - 24 * (len(worktime.index)) + worktime['WORKTIME'].sum() + bothtime / 60

    font = {'weight': 'bold',
            'size': 18}

    plt.rc('font', **font)

    plt.figure(figsize=(18, 10))
    plt.xlabel('Дата', fontsize=40)
    plt.ylabel('Время работы, ч', fontsize=40)
    plt.grid()
    plt.minorticks_on()
    plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
    plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
    plt.grid(which='minor',
             color='k',
             linestyle=':')
    plt.ylim([0, 25])
    plt.xlim([a, b])
    plt.xticks([i.date() for i in list(worktime['DATE'])[::4]], [i.date() for i in list(worktime['DATE'])[::4]])
    plt.plot(worktime['DATE'], worktime['WORKTIME'], label='1-й кластер', marker='s', markersize=15, color='darkblue',
             linewidth='6')
    box_1 = {'facecolor': 'white',  # цвет области
             'edgecolor': 'red',  # цвет крайней линии
             'boxstyle': 'round'}
    plt.title("1-кластер", bbox=box_1, fontsize=20, loc='center')
    plt.savefig('{}\\{}\\1worktime{}-{}-{}-{}.png'.format(pathpic, styear, stday, stmonth, endday, endmonth),
                bbox_inches='tight')
    time1path = "{}\\{}\\1worktime{}-{}-{}-{}.png".format(pathpic, styear, stday, stmonth, endday, endmonth)

    plt.figure(figsize=(18, 10))
    plt.xlabel('Дата', fontsize=40)
    plt.ylabel('Время работы, ч', fontsize=40)
    plt.grid()
    plt.minorticks_on()
    plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
    plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
    plt.ylim([0, 25])
    plt.xlim([a, b])
    plt.grid(which='minor', linestyle=':', color='k')
    plt.xticks([i.date() for i in list(worktime_2['DATE'])[::4]], [i.date() for i in list(worktime_2['DATE'])[::4]])
    plt.plot(worktime_2['DATE'], worktime_2['WORKTIME'], label='2-й кластер', marker='s', markersize=15,
             color='darkblue',
             linewidth='6')
    box_1 = {'facecolor': 'white',  # цвет области
             'edgecolor': 'red',  # цвет крайней линии
             'boxstyle': 'round'}
    plt.title("2-кластер", bbox=box_1, fontsize=20, loc='center')
    plt.savefig('{}\\{}\\2worktime{}-{}-{}-{}.png'.format(pathpic, styear, stday, stmonth, endday, endmonth),
                bbox_inches='tight')
    time2path = "{}\\{}\\2worktime{}-{}-{}-{}.png".format(pathpic, styear, stday, stmonth, endday, endmonth)

    a_4 = a4fr('', stday, endday, styear, endyear, stmonth, endmonth, file1cl)
    a_4_2 = a4fr(2, stday, endday, styear, endyear, stmonth, endmonth, file2cl)
    plt.figure(figsize=(18, 10))
    plt.xlabel('Дата', fontsize=40)
    plt.ylabel('N, соб/час', fontsize=40)
    plt.grid()
    plt.ylim([0, 5])
    plt.xlim([a, b])
    plt.minorticks_on()
    plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
    plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
    plt.grid(which='minor',
             color='k',
             linestyle=':')
    plt.xticks([i.date() for i in list(a_4['DATE'])[::4]], [i.date() for i in list(a_4['DATE'])[::4]])
    plt.plot(a_4['DATE'], a_4['EVENTS'] / worktime['WORKTIME'], label='1 Кл.', marker='s', markersize=15,
             color='darkblue', linewidth='6')
    plt.plot(a_4_2['DATE'], a_4_2['EVENTS'] / worktime_2['WORKTIME'], label='2 Кл.', marker='s', markersize=15,
             color='crimson', linewidth='6')
    plt.legend(bbox_to_anchor=(1.04, 1), loc="upper left")
    plt.savefig('{}\\{}\\amorethan54{}-{}-{}-{}.png'.format(pathpic, styear, stday, stmonth, endday, endmonth),
                bbox_inches='tight')
    a4f21path = "{}\\{}\\amorethan54{}-{}-{}-{}.png".format(pathpic, styear, stday, stmonth, endday, endmonth)

    nto0tr = ntozerotr('', stday, endday, styear, endyear, stmonth, endmonth, file1cl)

    nto0tr_2 = ntozerotr(2, stday, endday, styear, endyear, stmonth, endmonth, file2cl)
    # nto0tr, nto0tr_2 = change_name(start_date=date(styear, stmonth, stday), end_date=date(endyear, endmonth, endday))
    infonto0tr = nto0tr.describe().tail(7).head(2)
    infonto0tr_2 = nto0tr_2.describe().tail(7).head(2)
    infonto0tr.index = ['mean(100/cоб.)', 'std(100/cоб.)']
    infonto0tr_2.index = ['mean(100/cоб.)', 'std(100/cоб.)']

    plt.rc('axes',
           prop_cycle=(cycler('color', ['r', 'g', 'b', 'darkblue', 'lawngreen', 'hotpink', 'c', 'y', 'm', 'orange',
                                        'burlywood', 'darkmagenta', 'grey', 'darkslategray', 'saddlebrown',
                                        'lightsalmon'])))
    plt.figure(figsize=(18, 10))
    plt.xlabel('Дата', fontsize=40)
    plt.ylabel(r'$(coб)^{-1}$', fontsize=40)
    plt.grid()
    plt.minorticks_on()
    plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
    plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
    plt.grid(which='minor',
             color='k',
             linestyle=':')
    plt.ylim([0, 0.5], emit=False)
    plt.xlim([a, b])
    plt.xticks([i.date() for i in list(nto0tr['DATE'])[::4]], [i.date() for i in list(nto0tr['DATE'])[::4]])
    box_1 = {'facecolor': 'white',  # цвет области
             'edgecolor': 'red',  # цвет крайней линии
             'boxstyle': 'round'}
    plt.title("1-кластер", bbox=box_1, fontsize=20, loc='center')
    for i in range(1, 17):
        plt.plot(nto0tr['DATE'], nto0tr['n%s' % i], label='%s' % i, linewidth=5)
    plt.legend(bbox_to_anchor=(1.04, 1), loc="upper left")
    plt.savefig('{}\\{}\\ntozeromas{}-{}-{}-{}.png'.format(pathpic, styear, stday, stmonth, endday, endmonth),
                bbox_inches='tight')
    nzm1path = "{}\\{}\\ntozeromas{}-{}-{}-{}.png".format(pathpic, styear, stday, stmonth, endday, endmonth)

    plt.figure(figsize=(18, 10))
    plt.xlabel('Дата', fontsize=40)
    plt.ylabel(r'$(coб)^{-1}$', fontsize=40)
    plt.grid()
    plt.minorticks_on()
    plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
    plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
    plt.grid(which='minor',
             color='k',
             linestyle=':')
    plt.ylim([0, 0.5], emit=False)
    plt.xlim([a, b])
    plt.xticks([i.date() for i in list(nto0tr_2['DATE'])[::4]], [i.date() for i in list(nto0tr_2['DATE'])[::4]])
    box_1 = {'facecolor': 'white',  # цвет области
             'edgecolor': 'red',  # цвет крайней линии
             'boxstyle': 'round'}
    plt.title("2-кластер", bbox=box_1, fontsize=20, loc='center')
    for i in range(1, 17):
        plt.plot(nto0tr_2['DATE'], nto0tr_2['n%s' % i], label='%s' % i, linewidth=5)
    plt.legend(bbox_to_anchor=(1.04, 1), loc="upper left")
    plt.savefig('{}\\{}\\2ntozeromas{}-{}-{}-{}.png'.format(pathpic, styear, styear, stday, stmonth, endday, endmonth),
                bbox_inches='tight')
    nzm2path = "{}\\{}\\2ntozeromas{}-{}-{}-{}.png".format(pathpic, styear, styear, stday, stmonth, endday, endmonth)

    data = a52det('', stday, endday, styear, endyear, stmonth, endmonth, file1cl, amp=amp, fr=fr)
    data_2 = a52det(2, stday, endday, styear, endyear, stmonth, endmonth, file2cl, amp=amp, fr=fr)
    data_other = a52det('', stday, endday, styear, endyear, stmonth, endmonth, file1cl, amp=amp_2, fr=fr_2)
    data_other_2 = a52det(2, stday, endday, styear, endyear, stmonth, endmonth, file2cl, amp=amp_2, fr=fr_2)
    # data, data_2 = main(start_date=date(styear, stmonth, stday), end_date=date(endyear, endmonth, endday), amp=amp,
    #                     fr=fr)
    # data_other, data_other_2 = main(start_date=date(styear, stmonth, stday), end_date=date(endyear, endmonth, endday),
    #                                 amp=amp_2, fr=fr_2)
    plt.figure(figsize=(18, 10))
    plt.xlabel('Амплитуда, код АЦП', fontsize=40)
    plt.yscale('log')
    plt.xscale('log')
    plt.ylabel('Nсоб(Fr≥2, A>5)', fontsize=40)
    plt.minorticks_on()
    plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
    plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
    plt.xlim([amp - 1, 1000])
    plt.ylim([1, 1000])
    box_1 = {'facecolor': 'white',  # цвет области
             'edgecolor': 'red',  # цвет крайней линии
             'boxstyle': 'round'}
    plt.text(500, 500, "1-кластер", bbox=box_1, fontsize=20)
    for i in range(1, 17):
        plt.plot(data[data['e%s' % i] >= amp]['e%s' % i].value_counts().sort_index().keys().tolist(),
                 data[data['e%s' % i] >= amp]['e%s' % i].value_counts().sort_index(), label='%s' % i, linewidth=5)
    plt.legend(bbox_to_anchor=(1.04, 1), loc="upper left")
    plt.savefig(f'{pathpic}\\{styear}\\Neva{amp}fr{fr}{stday}-{stmonth}-{endday}-{endmonth}.png',
                bbox_inches='tight')
    fa25path = f'{pathpic}\\{styear}\\Neva{amp}fr{fr}{stday}-{stmonth}-{endday}-{endmonth}.png'

    plt.figure(figsize=(18, 10))
    plt.xlabel('Амплитуда, код АЦП', fontsize=40)
    plt.yscale('log')
    plt.xscale('log')
    plt.ylabel('Nсоб(Fr≥2, A>5)', fontsize=40)
    plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
    plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
    plt.xlim([amp - 1, 1000])
    plt.ylim([1, 1000])
    box_1 = {'facecolor': 'white',  # цвет области
             'edgecolor': 'red',  # цвет крайней линии
             'boxstyle': 'round'}
    plt.text(500, 500, "2-кластер", bbox=box_1, fontsize=20)
    for i in range(1, 17):
        plt.plot(data_2[data_2['e%s' % i] >= amp]['e%s' % i].value_counts().sort_index().keys().tolist(),
                 data_2[data_2['e%s' % i] >= amp]['e%s' % i].value_counts().sort_index(), label='%s' % i, linewidth=5)
    plt.legend(bbox_to_anchor=(1.04, 1), loc="upper left")
    plt.savefig(f'{pathpic}\\{styear}\\N2eva{amp}fr{fr}{stday}-{stmonth}-{endday}-{endmonth}.png',
                bbox_inches='tight')
    fa252path = f'{pathpic}\\{styear}\\N2eva{amp}fr{fr}{stday}-{stmonth}-{endday}-{endmonth}.png'

    plt.figure(figsize=(18, 10))
    plt.xlabel('Дата', fontsize=40)
    plt.ylabel('N, соб/час', fontsize=40)
    plt.grid()
    plt.minorticks_on()
    plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
    plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
    plt.xlim([a, b])
    plt.ylim([0, 8])
    # plt.rc('axes', prop_cycle=(cycler('color', ['r', 'g', 'b', 'darkblue', 'lawngreen', 'b', 'c', 'y', 'm', 'orange',
    #                                             'burlywood', 'darkmagenta', 'grey', 'darkslategray', 'saddlebrown',
    #                                             'purple'])))
    plt.grid(which='minor',
             color='k',
             linestyle=':')
    box_1 = {'facecolor': 'white',  # цвет области
             'edgecolor': 'red',  # цвет крайней линии
             'boxstyle': 'round'}
    plt.title("1-кластер", bbox=box_1, fontsize=20, loc='center')
    for i in range(1, 17):
        test = data[data['e%s' % i] > amp]['DATE'].value_counts()
        test = test.reindex(
            pd.to_datetime(pd.date_range(date(styear, stmonth, stday), date(endyear, endmonth, endday))),
            fill_value=0).sort_index().to_frame()
        test['VALUE'] = test['DATE']
        test['DATE'] = test.index
        test = test.merge(worktime, how='left')
        plt.xticks([i.date() for i in list(test['DATE'])[::4]], [i.date() for i in list(test['DATE'])[::4]])
        plt.plot(test['DATE'], test['VALUE'] / test['WORKTIME'], label='%s' % i, linewidth=6)

    plt.legend(bbox_to_anchor=(1.04, 1), loc="upper left")
    plt.savefig(f'{pathpic}\\{styear}\\va{amp}fr{fr}{stday}-{stmonth}-{endday}-{endmonth}.png',
                bbox_inches='tight')
    a2f5path = f'{pathpic}\\{styear}\\va{amp}fr{fr}{stday}-{stmonth}-{endday}-{endmonth}.png'

    plt.figure(figsize=(18, 10))
    plt.xlabel('Дата', fontsize=40)
    plt.ylabel('N, соб/час', fontsize=40)
    plt.grid()
    plt.minorticks_on()
    plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
    plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
    plt.xlim([a, b])
    plt.ylim([0, 8])
    plt.grid(which='minor',
             color='k',
             linestyle=':')
    box_1 = {'facecolor': 'white',  # цвет области
             'edgecolor': 'red',  # цвет крайней линии
             'boxstyle': 'round'}
    plt.title("2-кластер", bbox=box_1, fontsize=20, loc='center')
    for i in range(1, 17):
        test_2 = data_2.loc[data_2['e%s' % i] > amp, 'DATE'].value_counts()
        # print(test_2)
        # print(data_2['DATE'].unique())
        test_2 = test_2.reindex(
            pd.to_datetime(pd.date_range(date(styear, stmonth, stday), date(endyear, endmonth, endday))),
            fill_value=0).sort_index().to_frame()
        test_2['VALUE'] = test_2['DATE']
        test_2['DATE'] = test_2.index
        print(test_2)
        test_2 = test_2.merge(worktime_2, how='left')
        plt.xticks([i.date() for i in list(test_2['DATE'])[::4]], [i.date() for i in list(test_2['DATE'])[::4]])
        plt.plot(test_2['DATE'], test_2['VALUE'] / test_2['WORKTIME'], label='%s' % i, linewidth=6)

    plt.legend(bbox_to_anchor=(1.04, 1), loc="upper left")
    plt.savefig(f'{pathpic}\\{styear}\\v2a{amp}fr{fr}{stday}-{stmonth}-{endday}-{endmonth}.png',
                bbox_inches='tight')
    a2f52path = f'{pathpic}\\{styear}\\v2a{amp}fr{fr}{stday}-{stmonth}-{endday}-{endmonth}.png'

    plt.figure(figsize=(18, 10))
    plt.xlabel('Дата', fontsize=40)
    plt.ylabel('N, соб/час', fontsize=40)
    plt.grid()
    plt.minorticks_on()
    plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
    plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
    plt.xlim([a, b])
    plt.ylim([0, 8])
    # plt.rc('axes', prop_cycle=(cycler('color', ['r', 'g', 'b', 'darkblue', 'lawngreen', 'b', 'c', 'y', 'm', 'orange',
    #                                             'burlywood', 'darkmagenta', 'grey', 'darkslategray', 'saddlebrown',
    #                                             'purple'])))
    plt.grid(which='minor',
             color='k',
             linestyle=':')
    box_1 = {'facecolor': 'white',  # цвет области
             'edgecolor': 'red',  # цвет крайней линии
             'boxstyle': 'round'}
    plt.title("1-кластер", bbox=box_1, fontsize=20, loc='center')
    for i in range(1, 17):
        test = data_other[data_other['e%s' % i] > amp_2]['DATE'].value_counts()
        test = test.reindex(
            pd.to_datetime(pd.date_range(date(styear, stmonth, stday), date(endyear, endmonth, endday))),
            fill_value=0).sort_index().to_frame()
        test['VALUE'] = test['DATE']
        test['DATE'] = test.index
        test = test.merge(worktime, how='left')
        plt.xticks([i.date() for i in list(test['DATE'])[::4]], [i.date() for i in list(test['DATE'])[::4]])
        plt.plot(test['DATE'], test['VALUE'] / test['WORKTIME'], label='%s' % i, linewidth=6)
    plt.legend(bbox_to_anchor=(1.04, 1), loc="upper left")
    plt.savefig(f'{pathpic}\\{styear}\\va{amp_2}fr{fr_2}{stday}-{stmonth}-{endday}-{endmonth}.png',
                bbox_inches='tight')
    a1f10path = f'{pathpic}\\{styear}\\va{amp_2}fr{fr_2}{stday}-{stmonth}-{endday}-{endmonth}.png'

    plt.figure(figsize=(18, 10))
    plt.xlabel('Дата', fontsize=40)
    plt.ylabel('N, соб/час', fontsize=40)
    plt.grid()
    plt.minorticks_on()
    plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
    plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
    plt.xlim([a, b])
    plt.ylim([0, 8])
    # plt.rc('axes', prop_cycle=(cycler('color', ['r', 'g', 'b', 'darkblue', 'lawngreen', 'b', 'c', 'y', 'm', 'orange',
    #                                             'burlywood', 'darkmagenta', 'grey', 'darkslategray', 'saddlebrown',
    #                                             'purple'])))
    plt.grid(which='minor',
             color='k',
             linestyle=':')
    box_1 = {'facecolor': 'white',  # цвет области
             'edgecolor': 'red',  # цвет крайней линии
             'boxstyle': 'round'}
    plt.title("2-кластер", bbox=box_1, fontsize=20, loc='center')
    for i in range(1, 17):
        test = data_other_2[data_other_2['e%s' % i] > amp_2]['DATE'].value_counts()
        test = test.reindex(
            pd.to_datetime(pd.date_range(date(styear, stmonth, stday), date(endyear, endmonth, endday))),
            fill_value=0).sort_index().to_frame()
        test['VALUE'] = test['DATE']
        test['DATE'] = test.index
        test = test.merge(worktime_2, how='left')
        plt.xticks([i.date() for i in list(test['DATE'])[::4]], [i.date() for i in list(test['DATE'])[::4]])
        plt.plot(test['DATE'], test['VALUE'] / test['WORKTIME'], label='%s' % i, linewidth=6)
    plt.legend(bbox_to_anchor=(1.04, 1), loc="upper left")
    plt.savefig(f'{pathpic}\\{styear}\\v2a{amp_2}fr{fr_2}{stday}-{stmonth}-{endday}-{endmonth}.png',
                bbox_inches='tight')
    a1f102path = f'{pathpic}\\{styear}\\v2a{amp_2}fr{fr_2}{stday}-{stmonth}-{endday}-{endmonth}.png'

    df = pd.DataFrame()
    df_2 = pd.DataFrame()
    for i in range(1, 17):
        test = data[data['e%s' % i] > amp]['DATE'].value_counts()
        test = test.reindex(
            pd.to_datetime(pd.date_range(date(styear, stmonth, stday), date(endyear, endmonth, endday))),
            fill_value=0).sort_index().to_frame()
        test['VALUE'] = test['DATE']
        test['DATE'] = test.index
        test = test.merge(worktime, how='left')
        df["%s" % i] = test['VALUE'] / worktime[
            'WORKTIME']
    df = df.describe().tail(7).head(2)
    for j in range(1, 17):
        test_2 = data_2[data_2['e%s' % j] > amp]['DATE'].value_counts()
        test_2 = test_2.reindex(
            pd.to_datetime(pd.date_range(date(styear, stmonth, stday), date(endyear, endmonth, endday))),
            fill_value=0).sort_index().to_frame()
        test_2['VALUE'] = test_2['DATE']
        test_2['DATE'] = test_2.index
        test_2 = test_2.merge(worktime_2, how='left')
        df_2["%s" % j] = test_2['VALUE'] / worktime_2[
            'WORKTIME']
    df_2 = df_2.describe().tail(7).head(2)
    df.index = ['mean(соб./ч.)', 'std(соб./ч.)']
    df_2.index = ['mean(соб./ч.)', 'std(соб./ч.)']

    df_other = pd.DataFrame()
    df_other_2 = pd.DataFrame()
    for i in range(1, 17):
        test = data_other[data_other['e%s' % i] > amp_2]['DATE'].value_counts()
        test = test.reindex(
            pd.to_datetime(pd.date_range(date(styear, stmonth, stday), date(endyear, endmonth, endday))),
            fill_value=0).sort_index().to_frame()
        test['VALUE'] = test['DATE']
        test['DATE'] = test.index
        test = test.merge(worktime, how='left')
        df_other["%s" % i] = test['VALUE'] / worktime[
            'WORKTIME']
    df_other = df_other.describe().tail(7).head(2)
    for j in range(1, 17):
        test_2 = data_other_2[data_other_2['e%s' % j] > amp_2]['DATE'].value_counts()
        test_2 = test_2.reindex(
            pd.to_datetime(pd.date_range(date(styear, stmonth, stday), date(endyear, endmonth, endday))),
            fill_value=0).sort_index().to_frame()
        test_2['VALUE'] = test_2['DATE']
        test_2['DATE'] = test_2.index
        test_2 = test_2.merge(worktime_2, how='left')
        df_other_2["%s" % j] = test_2['VALUE'] / worktime_2[
            'WORKTIME']
    df_other_2 = df_other_2.describe().tail(7).head(2)
    df_other.index = ['mean(соб./ч.)', 'std(соб./ч.)']
    df_other_2.index = ['mean(соб./ч.)', 'std(соб./ч.)']

    doc = Document()

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    styles = doc.styles
    style = styles.add_style('PItalic', WD_STYLE_TYPE.PARAGRAPH)
    style = doc.styles['PItalic']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    font.italic = True

    headstyle = styles.add_style('Headstyle', WD_STYLE_TYPE.PARAGRAPH)
    headstyle = doc.styles['Headstyle']
    font = headstyle.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True

    headgraf = styles.add_style('Headgraf', WD_STYLE_TYPE.PARAGRAPH)
    headgraf = doc.styles['Headgraf']
    font = headgraf.font
    font.name = 'Times New Roman'
    font.size = Pt(13)
    font.bold = True
    font.italic = True

    head = doc.add_paragraph(
        'Справка о работе установки ПРИЗМА-32 в период с {:02}.{:02}.{} по {:02}.{:02}.{} '.format(stday, stmonth,
                                                                                                   styear,
                                                                                                   endday, endmonth,
                                                                                                   endyear),
        style='Headstyle')
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc = doc.add_paragraph('Таблица 1: Время работы установки ПРИЗМА-32.', style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    beg = doc.add_table(4, 4, doc.styles['Table Grid'])
    beg.cell(0, 0).text = '№ кластера'
    beg.cell(0, 1).text = 'Экспозиции, ч.'
    beg.cell(0, 2).text = 'Календарное время, ч.'
    beg.cell(0, 3).text = 'Экспозиция, %'
    beg.cell(1, 0).text = '1'
    beg.cell(1, 1).text = str(round(worktime['WORKTIME'].sum(), 2))
    beg.cell(1, 2).text = str(24 * len(pd.date_range(date(styear, stmonth, stday), date(endyear, endmonth, endday))))
    beg.cell(1, 3).text = str(round(worktime['WORKTIME'].sum() / (24 * (len(worktime))) * 100, 3)) + '%'
    beg.cell(2, 0).text = '2'
    beg.cell(2, 1).text = str(round(worktime_2['WORKTIME'].sum(), 2))
    beg.cell(2, 2).text = str(24 * len(pd.date_range(date(styear, stmonth, stday), date(endyear, endmonth, endday))))
    beg.cell(2, 3).text = str(round(worktime_2['WORKTIME'].sum() / (24 * (len(worktime_2))) * 100, 3)) + '%'
    beg.cell(3, 0).text = '1&2'
    beg.cell(3, 1).text = str(round(realtime, 2))
    beg.cell(3, 2).text = str(24 * len(pd.date_range(date(styear, stmonth, stday), date(endyear, endmonth, endday))))
    beg.cell(3, 3).text = str(round(realtime / (24 * (len(worktime))) * 100, 3)) + '%'

    for row in range(1):
        for col in range(4):
            # получаем ячейку таблицы
            cell = beg.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True

    for row in range(1, 4):
        for col in range(1):
            # получаем ячейку таблицы
            cell = beg.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
    space = doc.add_paragraph()
    desc = doc.add_paragraph('Таблица 2: Сводная таблица остановок и работ установки ПРИЗМА-32.', style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    err = doc.add_table(len(failstr_begin) + len(failstr_2_begin) + 2, 5, doc.styles['Table Grid'])
    err.alignment = WD_TABLE_ALIGNMENT.CENTER
    err.cell(0, 0).text = '№ кластера'
    err.cell(0, 0).merge(err.cell(1, 0))
    err.cell(0, 1).text = 'Время простоя'
    err.cell(1, 1).text = 'c'
    err.cell(1, 2).text = 'по'
    err.cell(0, 1).merge(err.cell(0, 2))
    err.cell(0, 3).text = 'Кол-во потерянных минут (период)'
    err.cell(0, 3).merge(err.cell(1, 3))
    err.cell(0, 4).text = 'Примечание'
    err.cell(0, 4).merge(err.cell(1, 4))

    for i in range(2, len(failstr_begin) + 2):
        err.cell(i, 0).text = '1'
        err.cell(i, 1).text = str(failstr_begin[i - 2])
        err.cell(i, 2).text = str(failstr_end[i - 2])
        err.cell(i, 3).text = str(lost_minutes[i - 2])
        err.cell(i, 4).text = ' '

    for i in range(2 + len(failstr_begin), len(failstr_2_begin) + 2 + len(failstr_begin)):
        err.cell(i, 0).text = '2'
        err.cell(i, 1).text = str(failstr_2_begin[i - 2 - len(failstr_begin)])
        err.cell(i, 2).text = str(failstr_2_end[i - 2 - len(failstr_begin)])
        err.cell(i, 3).text = str(lost_minutes_2[i - 2 - len(failstr_begin)])
        err.cell(i, 4).text = ' '

    for row in range(2):
        for col in range(5):
            # получаем ячейку таблицы
            cell = err.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True

    for row in range(1, len(failstr_begin) + 2 + len(failstr_2_begin)):
        for col in range(1):
            # получаем ячейку таблицы
            cell = err.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True

    for row in range(len(failstr_begin) + 2 + len(failstr_2_begin)):
        for col in range(5):
            cell = err.cell(row, col)
            # записываем в ячейку данные
            para_ph = cell.paragraphs[0]
            para_ph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # err = doc.add_table(3, 4, doc.styles['Table Grid'])
    # err.cell(0, 0).text = '№ кластера'
    # err.cell(0, 1).text = 'Время, простоя'
    # err.cell(0, 2).text = 'Кол-во остановок'
    # err.cell(0, 3).text = 'Причины остановок, описание поломок'
    # err.cell(1, 0).text = '1'
    # # err.cell(1, 1).text = failstr
    # err.cell(1, 2).text = str(breaks)
    # # err.cell(2,3).text=str(round(time_2['Unnamed: 1'].sum()/(24*(len(time_2)+1)),3))
    # err.cell(2, 0).text = '2'
    # # err.cell(2, 1).text = failstr_2
    # err.cell(2, 2).text = str(breaks_2)

    space = doc.add_paragraph()

    desc = doc.add_paragraph(
        'Таблица 3: Сводная таблица темпов счета событий и сигналов, отобранных как нейтрон кластеров установки ПРИЗМА-32.',
        style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    neut = doc.add_table(3, 3, doc.styles['Table Grid'])
    neut.cell(0, 0).text = 'Счет/кластер'
    neut.cell(0, 1).text = 'Кл1'
    neut.cell(0, 2).text = 'Кл2'
    neut.cell(1, 0).text = 'События (Fr ≥ 4, A ≥ 5), N соб./ч.'
    neut.cell(1, 1).text = str(round((a_4['EVENTS'] / worktime['WORKTIME']).mean(), 2))
    neut.cell(1, 2).text = str(round((a_4_2['EVENTS'] / worktime_2['WORKTIME']).mean(), 2))
    neut.cell(2, 0).text = 'Нейтроны, (Nn)/соб.'
    neut.cell(2, 1).text = str(round(infonto0tr.iloc[0].sum(), 2))
    neut.cell(2, 2).text = str(round(infonto0tr_2.iloc[0].sum(), 2))

    for row in range(1):
        for col in range(3):
            # получаем ячейку таблицы
            cell = neut.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True

    for row in range(1, 3):
        for col in range(1):
            # получаем ячейку таблицы
            cell = neut.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
    for cell in neut.columns[0].cells:
        cell.width = Inches(2.5)
    for cell in neut.columns[1].cells:
        cell.width = Inches(0.5)
    for cell in neut.columns[2].cells:
        cell.width = Inches(0.5)
    neut.alignment = WD_TABLE_ALIGNMENT.CENTER
    space = doc.add_paragraph()

    prim1 = doc.add_paragraph('')
    prim1.add_run('Примечание:').bold = True
    prim2 = doc.add_paragraph(
        '        В таблице 4 представлена сводная информация о неисправностях в работе детекторов кластера.')
    prim3 = doc.add_paragraph('Таблица 4: Неисправности.', style='PItalic')
    prim3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    clprim = doc.add_table(3, 5, doc.styles['Table Grid'])
    clprim.cell(0, 0).text = '№'
    clprim.cell(0, 1).text = 'Кластер'
    clprim.cell(0, 2).text = '№ Детектора'
    clprim.cell(0, 3).text = 'Период'
    clprim.cell(0, 4).text = 'Примечание'
    clprim.cell(1, 0).text = '1'
    clprim.cell(2, 0).text = '2'

    # space = doc.add_paragraph()
    # baddet = doc.add_paragraph('')
    # baddet.add_run('Наиболее плохо работающие детекторы (').bold = True
    # run = baddet.add_run('по приоритету')
    # font = run.font
    # font.color.rgb = RGBColor(217, 26, 26)
    # baddet.add_run(')').bold = True

    for cell in clprim.columns[0].cells:
        cell.width = Inches(0.3)
    for cell in clprim.columns[1].cells:
        cell.width = Inches(0.8)
    for cell in clprim.columns[2].cells:
        cell.width = Inches(1.2)
    for cell in clprim.columns[3].cells:
        cell.width = Inches(1)
    for cell in clprim.columns[4].cells:
        cell.width = Inches(4.2)
    neut.alignment = WD_TABLE_ALIGNMENT.CENTER

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    desc = doc.add_paragraph('Продолжительность работы кластеров установки ПРИЗМА-32.', style='Headgraf')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(time1path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc = doc.add_paragraph('Рис. 1 - Продолжительность работы 1-го кластера в сутки', style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(time2path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc = doc.add_paragraph('Рис. 2 - Продолжительность работы 2-го кластера в сутки', style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    desc = doc.add_paragraph('Скорость счета событий установки ПРИЗМА-32', style='Headgraf')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(a4f21path, width=Inches(7))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc = doc.add_paragraph('Рис. 3 - Скорость счета событий Fr ≥ 4, A ≥ 5', style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    doc.add_picture(nzm1path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc = doc.add_paragraph('Рис. 4 - Число импульсов в событии, отобранных как нейтрон, при самозапуске кластер 1',
                             style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_picture(nzm2path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc = doc.add_paragraph('Рис. 5 - Число импульсов в событии, отобранных как нейтрон, при самозапуске кластер 2',
                             style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    desc = doc.add_paragraph(
        'Таблица 5: Среднее число нейтронов (Nn) для детекторов установки ПРИЗМА-32 за месяц работы, нормированное на количество событий (Ns).(при самозапуске), ',
        style='PItalic')
    desc.add_run('(100/соб)').bold = True
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    b = doc.add_table(infonto0tr.shape[0] + infonto0tr_2.shape[0] + 2, infonto0tr.shape[1] + 2,
                      doc.styles['Table Grid'])

    b.cell(0, 0).text = "№"
    b.cell(0, 1).text = "Стат-ка"
    b.cell(0, 2).text = "№ детектора"
    b.cell(0, 2).merge(b.cell(0, 17))
    b.cell(0, 0).merge(b.cell(1, 0))
    b.cell(0, 1).merge(b.cell(1, 1))
    for j in range(infonto0tr.shape[-1]):
        b.cell(1, j + 2).text = infonto0tr.columns[j]
    b.cell(2, 0).text = '1'
    b.cell(2, 0).merge(b.cell(3, 0))
    b.cell(4, 0).text = '2'
    b.cell(4, 0).merge(b.cell(5, 0))
    for i in range(infonto0tr.shape[0]):
        b.cell(i + 2, 1).text = infonto0tr.index[i]
        for j in range(infonto0tr.shape[-1]):
            b.cell(i + 2, j + 2).text = str(round(infonto0tr.values[i, j], 2))

    for i in range(infonto0tr_2.shape[0]):
        b.cell(i + 2 + infonto0tr.shape[0], 1).text = infonto0tr_2.index[i]
        for j in range(infonto0tr_2.shape[-1]):
            b.cell(i + 2 + infonto0tr.shape[0], j + 2).text = str(round(infonto0tr_2.values[i, j], 2))

    for row in b.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(7.5)

    for row in range(2):
        for col in range(18):
            # получаем ячейку таблицы
            cell = b.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(9)

    for row in range(2, 6):
        for col in range(2):
            # получаем ячейку таблицы
            cell = b.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(8.5)

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    doc.add_picture(a2f5path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # desc=doc.add_paragraph('Рис. 9 - Скорость счета сигналов дететоров А ≥ 5 кластер 1', style='PItalic')
    desc = doc.add_paragraph(f'Рис. 6 - Скорость счета  детекторов в 1-м кластере Fr ≥ {fr}, A > {amp - 1}',
                             style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(a2f52path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # desc=doc.add_paragraph('Рис. 10 - Скорость счета сигналов детекторов А ≥ 5 кластер 2', style='PItalic')
    desc = doc.add_paragraph(f'Рис. 7 - Скорость счета детекторов во 2-м кластере Fr ≥ {fr}, A > {amp - 1}',
                             style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print(f'fr - {fr}, amp - {amp - 1}')
    desc = doc.add_paragraph('Таблица 6: Среднемесячные число срабатываний детекторов установки ПРИЗМА-32, cоб./час.',
                             style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    c = doc.add_table(df.shape[0] + df_2.shape[0] + 2, df.shape[1] + 2, doc.styles['Table Grid'])

    c.cell(0, 0).text = "№"
    c.cell(0, 1).text = "Стат-ка"
    c.cell(0, 2).text = "№ детектора"
    c.cell(0, 2).merge(c.cell(0, 17))
    c.cell(0, 0).merge(c.cell(1, 0))
    c.cell(0, 1).merge(c.cell(1, 1))
    for j in range(df.shape[-1]):
        c.cell(1, j + 2).text = df.columns[j]
    c.cell(2, 0).text = '1'
    c.cell(2, 0).merge(c.cell(3, 0))
    c.cell(4, 0).text = '2'
    c.cell(4, 0).merge(c.cell(5, 0))
    for i in range(df.shape[0]):
        c.cell(i + 2, 1).text = df.index[i]
        for j in range(df.shape[-1]):
            c.cell(i + 2, j + 2).text = str(round(df.values[i, j], 2))

    for i in range(df_2.shape[0]):
        c.cell(i + 2 + df.shape[0], 1).text = df_2.index[i]
        for j in range(df_2.shape[-1]):
            c.cell(i + 2 + df.shape[0], j + 2).text = str(round(df_2.values[i, j], 2))

    for row in c.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    for row in range(2):
        for col in range(18):
            # получаем ячейку таблицы
            cell = c.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(9)

    for row in range(2, 6):
        for col in range(2):
            # получаем ячейку таблицы
            cell = c.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(8.5)

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    doc.add_picture(a1f10path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # desc=doc.add_paragraph('Рис. 9 - Скорость счета сигналов дететоров А ≥ 5 кластер 1', style='PItalic')
    desc = doc.add_paragraph(f'Рис. 8 - Скорость счета  детекторов в 1-м кластере Fr ≥ {fr_2}, A > {amp_2 - 1}',
                             style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_picture(a1f102path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # desc=doc.add_paragraph('Рис. 10 - Скорость счета сигналов детекторов А ≥ 5 кластер 2', style='PItalic')
    desc = doc.add_paragraph(f'Рис. 9 - Скорость счета детекторов во 2-м кластере Fr ≥ {fr_2}, A > {amp_2 - 1}',
                             style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    print(f'fr - {fr}, amp - {amp - 1}')
    print(f'fr - {fr_2}, amp - {amp_2 - 1}')
    desc = doc.add_paragraph('Таблица 6: Среднемесячные число срабатываний детекторов установки ПРИЗМА-32, cоб./час.',
                             style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    g = doc.add_table(df_other.shape[0] + df_other_2.shape[0] + 2, df_other.shape[1] + 2, doc.styles['Table Grid'])

    g.cell(0, 0).text = "№"
    g.cell(0, 1).text = "Стат-ка"
    g.cell(0, 2).text = "№ детектора"
    g.cell(0, 2).merge(g.cell(0, 17))
    g.cell(0, 0).merge(g.cell(1, 0))
    g.cell(0, 1).merge(g.cell(1, 1))
    for j in range(df_other.shape[-1]):
        g.cell(1, j + 2).text = df_other.columns[j]
    g.cell(2, 0).text = '1'
    g.cell(2, 0).merge(g.cell(3, 0))
    g.cell(4, 0).text = '2'
    g.cell(4, 0).merge(g.cell(5, 0))
    for i in range(df_other.shape[0]):
        g.cell(i + 2, 1).text = df_other.index[i]
        for j in range(df_other.shape[-1]):
            g.cell(i + 2, j + 2).text = str(round(df_other.values[i, j], 2))

    for i in range(df_other_2.shape[0]):
        g.cell(i + 2 + df_other.shape[0], 1).text = df_other_2.index[i]
        for j in range(df_other_2.shape[-1]):
            g.cell(i + 2 + df_other.shape[0], j + 2).text = str(round(df_other_2.values[i, j], 2))

    for row in g.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)

    for row in range(2):
        for col in range(18):
            # получаем ячейку таблицы
            cell = g.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(9)

    for row in range(2, 6):
        for col in range(2):
            # получаем ячейку таблицы
            cell = g.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
            run.font.size = Pt(8.5)

    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)

    desc = doc.add_paragraph('На рисунке 8, 9 представлено число сигналов с А>5 кодов АЦП в час для 16 детекторов.',
                             style='Headgraf')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_picture(fa25path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    desc = doc.add_paragraph(
        f'Рис. 8 - Амплитудное распределение сигналов от детекторов  кластер 1 (Fr ≥ {fr} и А > {amp - 1})',
        style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_picture(fa252path, width=Inches(6))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    desc = doc.add_paragraph(
        f'Рис. 9 - Амплитудное распределение сигналов от детекторов  кластер 2 (Fr ≥ {fr} и А > {amp - 1})',
        style='PItalic')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_page_number(doc.sections[0].footer.paragraphs[0])

    doc.save(f'{path}\{stday:02}.{stmonth:02}.{styear}-{endday:02}-{endmonth:02}.{endyear}.docx')

    plt.close('all')

    # ДЛЯ ГОДОВОГО ОТЧЕТА
    # plt.figure(figsize=(18, 10))
    # plt.xlabel('Дата', fontsize=20)
    # plt.ylabel(r'$(coб)^{-1}$', fontsize=20)
    # plt.grid()
    # plt.minorticks_on()
    # plt.tick_params(axis='both', which='minor', direction='out', length=10, width=2, pad=10)
    # plt.tick_params(axis='both', which='major', direction='out', length=20, width=4, pad=10)
    # plt.grid(which='minor',
    #          color='k',
    #          linestyle=':')
    # plt.ylim([0, 0.5], emit=False)
    # plt.xlim([a, b])
    # box_1 = {'facecolor': 'white',  # цвет области
    #          'edgecolor': 'red',  # цвет крайней линии
    #          'boxstyle': 'round'}
    # plt.title("1-кластер", bbox=box_1, fontsize=20, loc='center')
    # for i in range(1, 17):
    #     plt.plot(nto0tr['DATE'], nto0tr['n%s' % i], label='%s' % i, linewidth=5)
    # plt.legend(bbox_to_anchor=(1.04, 1), loc="upper left")
    # plt.savefig('{}\\{}\\ntozeromas{}-{}-{}-{}.png'.format(pathpic, styear, stday, stmonth, endday, endmonth),
    #             bbox_inches='tight')
    # nzm1path = "{}\\{}\\ntozeromas{}-{}-{}-{}.png".format(pathpic, styear, stday, stmonth, endday, endmonth)
